from __future__ import annotations

import json
import re
from pathlib import Path
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config.settings import settings
from utils.hash_utils import scholarship_hash as build_scholarship_hash

class AutoApplyAgent:
    def __init__(self, output_dir: Path | None = None) -> None:
        self.output_dir = output_dir or (settings.DATA_DIR / "applications")
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate_postulation_package(
        self,
        *,
        profile: dict,
        scholarship: dict,
        package: dict,
        user_sub: str,
    ) -> dict:
        """
        Takes the reviewable application package and exports it as:
        1. A highly professional Word Motivation Letter (.docx)
        2. A standard Autofill JSON payload for browser automation (.json)
        """
        scholarship_name = scholarship.get("scholarship_name") or scholarship.get("name") or "Beca"
        source_url = scholarship.get("source_url") or scholarship.get("url") or ""
        
        # Build stable scholarship hash
        s_hash = build_scholarship_hash(scholarship_name, source_url)
        safe_sub = re.sub(r"[^a-zA-Z0-9_-]", "_", user_sub)
        
        # 1. Generate professional Word Document
        docx_filename = f"letter_{s_hash}_{safe_sub}.docx"
        docx_path = self.output_dir / docx_filename
        
        letter_content = package.get("letter_of_intent") or ""
        self._generate_docx(
            content=letter_content,
            path=docx_path,
            scholarship_name=scholarship_name,
            institution=scholarship.get("institution") or "",
        )
        
        # 2. Generate standard Autofill JSON payload
        json_filename = f"autofill_{s_hash}_{safe_sub}.json"
        json_path = self.output_dir / json_filename
        
        autofill_payload = {
            "metadata": {
                "scholarship_name": scholarship_name,
                "institution": scholarship.get("institution", ""),
                "application_url": package.get("application_url", ""),
                "generated_at": package.get("generated_at", ""),
                "user_sub": user_sub,
            },
            "fields": package.get("filled_fields", []),
            "checklist": package.get("submission_checklist", []),
        }
        
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(autofill_payload, f, indent=2, ensure_ascii=False)
            
        return {
            "docx_path": str(docx_path),
            "docx_filename": docx_filename,
            "json_path": str(json_path),
            "json_filename": json_filename,
            "autofill_payload": autofill_payload,
        }

    def _generate_docx(self, content: str, path: Path, scholarship_name: str, institution: str) -> None:
        doc = Document()
        
        # Set margins to 1 inch standard
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            
        # 1. Document Title / Header
        header_p = doc.add_paragraph()
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_p.add_run("CARTA DE INTENCIÓN / DECLARACIÓN DE MOTIVACIÓN")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(15, 23, 42) # Slate-900
        
        # Spacer
        doc.add_paragraph()
        
        # 2. Metadata Block (aligned left)
        meta_p = doc.add_paragraph()
        meta_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add bold metadata fields
        r = meta_p.add_run("Oportunidad: ")
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(10.5)
        
        r = meta_p.add_run(f"{scholarship_name}\n")
        r.font.name = "Arial"
        r.font.size = Pt(10.5)
        
        if institution:
            r = meta_p.add_run("Institución: ")
            r.bold = True
            r.font.name = "Arial"
            r.font.size = Pt(10.5)
            
            r = meta_p.add_run(f"{institution}\n")
            r.font.name = "Arial"
            r.font.size = Pt(10.5)
            
        r = meta_p.add_run("Generado por: ")
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(10.5)
        
        r = meta_p.add_run("DevIALabs Inteligente Multi-Agente\n")
        r.font.name = "Arial"
        r.font.size = Pt(10.5)
        
        # Horizontal line block
        border_p = doc.add_paragraph()
        border_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        border_p.add_run("-" * 75).font.color.rgb = RGBColor(226, 232, 240) # Slate-200
        
        # 3. Main Letter Body
        # Clean up double linebreaks for neat paragraph styling
        paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
        
        # Check if basic fallback notice is present, and format accordingly
        for p_text in paragraphs:
            body_p = doc.add_paragraph()
            body_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Paragraph formatting
            p_format = body_p.paragraph_format
            p_format.space_after = Pt(10)
            p_format.line_spacing = 1.15
            
            r = body_p.add_run(p_text)
            r.font.name = "Arial"
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(51, 65, 85) # Slate-700
            
        # Save file physically
        doc.save(path)
